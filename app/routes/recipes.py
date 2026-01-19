from fastapi import APIRouter, Request, Form
from fastapi.templating import Jinja2Templates
from fastapi.responses import RedirectResponse, FileResponse, PlainTextResponse
import sqlite3
import tempfile
from pathlib import Path
from typing import List
from recipes.recipe_database import (
    create_recipe,
    load_recipe,
    add_recipe_ingredient,
    update_recipe,
    replace_recipe_ingredients,
)
from ingredients.database import get_all_ingredient_names
from logic.recalculate_nutrients import recalculate_nutrients
from logic.scaled_recipe import scale_recipe_to_energy


from config import RECIPE_DB_PATH

router = APIRouter(prefix="/recipes")
templates = Jinja2Templates(directory="app/templates")


# Route to list all recipes from the database
@router.get("")
def list_recipes(request: Request):
    with sqlite3.connect(RECIPE_DB_PATH) as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, name, category
            FROM recipes
            ORDER BY category, name
        """
        )
        recipes = cur.fetchall()

    recipes_by_category = {
        "breakfast": [],
        "lunch": [],
        "dinner": [],
        "snack": [],
    }
    for recipe in recipes:
        recipes_by_category.setdefault(recipe[2], []).append(recipe)

    return templates.TemplateResponse(
        "recipes_list.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
        },
    )


# Route to show meal plan selection (export)
@router.get("/meal-plan")
def meal_plan(request: Request):
    with sqlite3.connect(RECIPE_DB_PATH) as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, name, category
            FROM recipes
            ORDER BY category, name
        """
        )
        recipes = cur.fetchall()

    recipes_by_category = {
        "breakfast": [],
        "lunch": [],
        "dinner": [],
        "snack": [],
    }
    for recipe in recipes:
        recipes_by_category.setdefault(recipe[2], []).append(recipe)

    return templates.TemplateResponse(
        "meal_plan.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
        },
    )


# Route to show the form for adding a new recipe
@router.get("/new")
def new_recipe_form(request: Request):
    return templates.TemplateResponse(
        "recipe_form.html",
        {"request": request},
    )


# Route to handle the submission of a new recipe
@router.post("/new")
def create_recipe_post(
    name: str = Form(...),
    category: str = Form(...),
    instructions: str = Form(...),
    image_path: str = Form(""),
):
    recipe_id = create_recipe(
        name=name,
        category=category,
        instructions=instructions,
        image_path=image_path,
    )

    return RedirectResponse(
        url=f"/recipes/{recipe_id}",
        status_code=303,
    )


@router.post("/export")
def export_recipes(request: Request, recipe_ids: str = Form("")):
    ids = []
    for rid in recipe_ids.split(","):
        rid = rid.strip()
        if not rid:
            continue
        try:
            ids.append(int(rid))
        except ValueError:
            continue
    recipes = []
    for rid in ids:
        recipe = load_recipe(rid)
        if not recipe:
            continue
        nutrients = recalculate_nutrients(recipe)
        recipes.append(
            {
                "id": rid,
                "energy_kj": nutrients.get("energy_kj", 0) or 0,
                **recipe,
            }
        )
    recipes_by_category = {
        "breakfast": [],
        "lunch": [],
        "dinner": [],
        "snack": [],
    }
    for recipe in recipes:
        recipes_by_category.setdefault(recipe["category"], []).append(recipe)

    return templates.TemplateResponse(
        "recipes_export.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
        },
    )


@router.post("/export/scale")
def scale_selected_recipes(
    request: Request,
    recipe_ids: List[int] = Form(...),
    client_name: str = Form(""),
    target_kj_breakfast: float = Form(0),
    target_kj_lunch: float = Form(0),
    target_kj_dinner: float = Form(0),
    target_kj_snack: float = Form(0),
):
    targets_by_category = {
        "breakfast": target_kj_breakfast,
        "lunch": target_kj_lunch,
        "dinner": target_kj_dinner,
        "snack": target_kj_snack,
    }
    recipes = []
    for rid in recipe_ids:
        recipe = load_recipe(rid)
        if recipe:
            recipes.append((rid, recipe))

    scaled_recipes = []
    for rid, recipe in recipes:
        category = recipe.get("category", "")
        target_kj = targets_by_category.get(category, 0) or 0
        if target_kj <= 0:
            continue
        scaled_recipe, scaled_nutrients = scale_recipe_to_energy(recipe, target_kj)
        scaled_recipes.append(
            {
                "recipe_id": rid,
                "recipe": scaled_recipe,
                "nutrients": scaled_nutrients,
                "category": category,
            }
        )
    return templates.TemplateResponse(
        "recipes_scaled_preview.html",
        {
            "request": request,
            "scaled_recipes": scaled_recipes,
            "targets_by_category": targets_by_category,
            "client_name": client_name,
            "error": None,
        },
    )


@router.post("/export/preview-update")
async def update_preview(request: Request):
    form = await request.form()
    recipe_ids = form.getlist("recipe_ids")
    client_name = form.get("client_name") or ""
    targets_by_category = {
        "breakfast": float(form.get("target_kj_breakfast") or 0),
        "lunch": float(form.get("target_kj_lunch") or 0),
        "dinner": float(form.get("target_kj_dinner") or 0),
        "snack": float(form.get("target_kj_snack") or 0),
    }

    scaled_recipes = []
    for rid_value in recipe_ids:
        try:
            rid = int(rid_value)
        except ValueError:
            continue
        recipe = load_recipe(rid)
        if not recipe:
            continue
        names = form.getlist(f"ingredient_name_{rid}")
        portions = form.getlist(f"portion_g_{rid}")
        ingredients = []
        for name, portion_value in zip(names, portions):
            try:
                portion = float(portion_value)
            except (TypeError, ValueError):
                continue
            ingredients.append({"name": name, "portion_g": portion})
        recipe["ingredients"] = ingredients
        nutrients = recalculate_nutrients(recipe)
        scaled_recipes.append(
            {
                "recipe_id": rid,
                "recipe": recipe,
                "nutrients": nutrients,
                "category": recipe.get("category", ""),
            }
        )

    return templates.TemplateResponse(
        "recipes_scaled_preview.html",
        {
            "request": request,
            "scaled_recipes": scaled_recipes,
            "targets_by_category": targets_by_category,
            "client_name": client_name,
            "error": None,
        },
    )


@router.post("/export/docx")
async def export_docx(request: Request):
    try:
        from docx import Document
    except ModuleNotFoundError:
        return PlainTextResponse(
            "python-docx is not installed. Run: pip install python-docx",
            status_code=500,
        )

    form = await request.form()
    recipe_ids = form.getlist("recipe_ids")
    client_name = (form.get("client_name") or "").strip()

    recipes = []
    for rid_value in recipe_ids:
        try:
            rid = int(rid_value)
        except ValueError:
            continue
        recipe = load_recipe(rid)
        if not recipe:
            continue
        names = form.getlist(f"ingredient_name_{rid}")
        portions = form.getlist(f"portion_g_{rid}")
        ingredients = []
        for name, portion_value in zip(names, portions):
            try:
                portion = float(portion_value)
            except (TypeError, ValueError):
                continue
            ingredients.append({"name": name, "portion_g": portion})
        recipe["ingredients"] = ingredients
        nutrients = recalculate_nutrients(recipe)
        recipes.append({"recipe": recipe, "nutrients": nutrients})

    doc = Document()
    doc.add_heading("NutriCoach Export", level=1)
    if client_name:
        doc.add_paragraph(f"Client: {client_name}")

    for item in recipes:
        recipe = item["recipe"]
        nutrients = item["nutrients"]
        doc.add_heading(recipe.get("name", "Recipe"), level=2)
        doc.add_paragraph(f"Category: {recipe.get('category', '')}")

        image_path = recipe.get("image_path")
        if image_path and Path(image_path).exists():
            try:
                doc.add_picture(image_path, width=None)
            except Exception:
                pass

        doc.add_paragraph("Ingredients:")
        for ing in recipe.get("ingredients", []):
            doc.add_paragraph(
                f"- {ing.get('name', '')} – {ing.get('portion_g', 0)} g",
                style="List Bullet",
            )

        instructions = recipe.get("instructions") or ""
        if instructions:
            doc.add_paragraph("Instructions:")
            doc.add_paragraph(instructions)

        doc.add_paragraph("Nutrition summary:")
        doc.add_paragraph(f"Energy: {nutrients['energy_kj']} kJ")
        doc.add_paragraph(f"Protein: {nutrients['protein_g']} g")
        doc.add_paragraph(f"Carbs: {nutrients['carbs_g']} g")
        doc.add_paragraph(f"Fat: {nutrients['fat_g']} g")
        doc.add_paragraph(f"Fibre: {nutrients['fibre_g']} g")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name
    doc.save(temp_path)

    return FileResponse(
        temp_path,
        filename="nutricoach_export.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


# Route to show the details of a specific recipe
@router.get("/{recipe_id}")
def recipe_detail(request: Request, recipe_id: int):
    recipe = load_recipe(recipe_id)
    ingredients = get_all_ingredient_names()

    nutrients = recalculate_nutrients(recipe)

    return templates.TemplateResponse(
        "recipe_detail.html",
        {
            "request": request,
            "recipe_id": recipe_id,
            "recipe": recipe,
            "ingredients": ingredients,
            "nutrients": nutrients,
        },
    )


# Route to handle adding an ingredient to a specific recipe
@router.post("/{recipe_id}")
def add_ingredient_to_recipe(
    recipe_id: int,
    ingredient_name: str = Form(...),
    portion_g: float = Form(...),
):
    add_recipe_ingredient(
        recipe_id=recipe_id,
        ingredient_name=ingredient_name,
        portion_g=portion_g,
    )

    return RedirectResponse(
        url=f"/recipes/{recipe_id}",
        status_code=303,
    )


@router.post("/{recipe_id}/edit")
def update_recipe_post(
    recipe_id: int,
    name: str = Form(...),
    category: str = Form(...),
    instructions: str = Form(...),
    image_path: str = Form(""),
    ingredient_name: List[str] = Form([]),
    portion_g: List[str] = Form([]),
):
    update_recipe(
        recipe_id=recipe_id,
        name=name,
        category=category,
        instructions=instructions,
        image_path=image_path,
    )
    ingredients = []
    for name_value, portion_value in zip(ingredient_name, portion_g):
        name_value = (name_value or "").strip()
        if not name_value:
            continue
        try:
            portion = float(portion_value)
        except (TypeError, ValueError):
            continue
        if portion <= 0:
            continue
        ingredients.append({"name": name_value, "portion_g": portion})
    replace_recipe_ingredients(recipe_id, ingredients)

    return RedirectResponse(
        url=f"/recipes/{recipe_id}",
        status_code=303,
    )
